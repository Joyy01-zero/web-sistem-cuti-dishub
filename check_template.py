"""Periksa isi template docx — cari teks nomor surat sebenarnya."""
import os

from docx import Document

TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template_surat", "template_cuti_pkwt.docx")
doc = Document(TEMPLATE)

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i}] {t!r}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            t = cell.text.strip()
            if t:
                print(f"T{ti}R{ri}C{ci}: {t!r}")
