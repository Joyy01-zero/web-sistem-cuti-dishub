import io
import os
import logging

from docx import Document
from docx.shared import Pt

from config.constants import BULAN_NAMA

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "template_surat",
    "template_cuti_pkwt.docx",
)


def _safe_str(val):
    """Convert any value to clean string."""
    if val is None:
        return ""
    return str(val).strip()


def generate_surat(data: dict) -> bytes:
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template surat tidak ditemukan: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)

    # Cast semua field ke str — Sheets kadang return int
    nama = _safe_str(data.get("NAMA", ""))
    jabatan = _safe_str(data.get("JABATAN", ""))
    seksi = _safe_str(data.get("SEKSI", ""))
    shif = _safe_str(data.get("SHIF", ""))
    hari = _safe_str(data.get("HARI", ""))
    keperluan = _safe_str(data.get("KEPERLUAN", ""))
    kabid = _safe_str(data.get("KABID/KASI", ""))
    no_surat = _safe_str(data.get("NO SURAT", ""))
    tgl_submit = _safe_str(data.get("TGL_SUBMIT", ""))

    logger.info(f"generate_surat: nama={nama}, no_surat={no_surat!r}")

    # Build replacement map (teks asli di template → nilai baru)
    text_replacements = {
        "WAHYU EKO SAPUTRO": nama,
        "Tenaga Operasional Lalu Lintas": jabatan,
        "Pengendalian dan Ketertiban": seksi,
        "28 s.d. 29 Juli 2026": hari,
        "KEPERLUAN KELUARGA": keperluan.upper() if keperluan else "",
        "FAIZAL RACHMAN, S.E": kabid,
    }

    # Shif
    shif_replacements = {"Normal": shif} if shif else {}

    # No surat — replace sesuai format
    no_surat_replacements = {}
    if no_surat:
        parts = no_surat.split("/")
        if len(parts) >= 4:
            nomor = parts[0].strip()
            bulan_romawi = parts[2].strip()
            tahun = parts[3].strip()
            no_surat_replacements = {
                "167": nomor,
                "VI": bulan_romawi,
            }
            # Hanya replace tahun surat, BUKAN semua "2026" di dokumen
            # Cari pola "167 /PKWT/VI/2026" sebagai string utuh dulu
            full_pattern = "167 /PKWT/VI/2026"
            no_surat_replacements[full_pattern] = f"{nomor} /PKWT/{bulan_romawi}/{tahun}"
        else:
            # Format tidak lengkap (misal cuma "121212") — ganti seluruh pola nomor template
            no_surat_replacements = {
                "167 /PKWT/VI/2026": no_surat,
                "167": no_surat,
            }

    # Tanggal surat (bagian 2)
    tgl_str = ""
    if tgl_submit:
        try:
            from datetime import datetime
            dt = datetime.strptime(tgl_submit[:10], "%Y-%m-%d")
            tgl_str = f"{BULAN_NAMA[dt.month]} {dt.year}"
        except Exception:
            tgl_str = tgl_submit

    date_replacements = {}
    if tgl_str:
        date_replacements["Juli 2026"] = tgl_str
        date_replacements["Maret 2026"] = tgl_str

    # Merge all
    all_replacements = {}
    all_replacements.update(text_replacements)
    all_replacements.update(shif_replacements)
    all_replacements.update(no_surat_replacements)
    all_replacements.update(date_replacements)

    # Replace in paragraphs (skip last empty paragraph to avoid blank page)
    all_paragraphs = doc.paragraphs
    for i, paragraph in enumerate(all_paragraphs):
        # Skip the final paragraph if empty (prevents extra blank page in Word)
        if i == len(all_paragraphs) - 1 and not paragraph.text.strip():
            # Set to minimal size so it doesn't force a new page
            for run in paragraph.runs:
                run.font.size = Pt(1)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            continue
        _replace_in_paragraph(paragraph, all_replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, all_replacements)

    # Remove trailing empty paragraphs to prevent blank page
    body = doc.element.body
    trailing_empty = []
    for child in reversed(list(body)):
        if child.tag.endswith('}p'):  # paragraph element
            # Collect all text from paragraph and its child runs
            p_text = ''.join(child.itertext()).strip()
            if not p_text:
                trailing_empty.append(child)
            else:
                break
        else:
            break
    for elem in trailing_empty:
        body.remove(elem)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _replace_in_paragraph(paragraph, replacements):
    """Replace text across runs, handling cases where target spans multiple runs."""
    full_text = paragraph.text
    if not any(old in full_text for old in replacements):
        return

    for old_text, new_text in replacements.items():
        if old_text not in full_text:
            continue

        # Pastikan new_text adalah string
        new_text = str(new_text) if new_text is not None else ""

        # Try simple run-level replacement first
        replaced = False
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced = True
                break

        if replaced:
            continue

        # Target spans multiple runs — concatenate, replace, redistribute
        runs = paragraph.runs
        if not runs:
            continue

        combined = "".join(r.text for r in runs)
        if old_text not in combined:
            continue

        new_combined = combined.replace(old_text, new_text)

        # Put all text in first run, clear rest
        runs[0].text = new_combined
        for r in runs[1:]:
            r.text = ""
