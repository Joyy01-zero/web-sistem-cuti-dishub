"""Run once to generate template_cuti_pkwt.docx with placeholders."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)


# ===== BAGIAN 1: SURAT PERMOHONAN CUTI =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SURAT PERMOHONAN CUTI")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph("Bogor, {{TANGGAL_SURAT}}")
doc.add_paragraph("")
doc.add_paragraph("Kepada Yth.")
doc.add_paragraph("Kepala Sub. Bag Umum dan Kepegawaian")
doc.add_paragraph("Dinas Perhubungan Kota Bogor")
doc.add_paragraph("di Tempat")
doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("Perihal: Permohonan Cuti")
run.bold = True

doc.add_paragraph("")
doc.add_paragraph("Dengan hormat,")
doc.add_paragraph(
    "Yang bertanda tangan di bawah ini, saya mengajukan permohonan cuti dengan data sebagai berikut:"
)
doc.add_paragraph("")

# Table data karyawan
table = doc.add_table(rows=8, cols=2)
table.style = "Table Grid"
data_rows = [
    ("Nama", "{{NAMA}}"),
    ("NIP", "{{NIP}}"),
    ("Jabatan", "{{JABATAN}}"),
    ("Bidang/Seksi", "{{SEKSI}}"),
    ("Shif", "{{SHIF}}"),
    ("Tanggal Cuti", "{{TANGGAL_CUTI}}"),
    ("Keperluan", "{{KEPERLUAN}}"),
    ("Kabid/Kasi", "{{KABID_KASI}}"),
]
for i, (label, placeholder) in enumerate(data_rows):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = placeholder

doc.add_paragraph("")
doc.add_paragraph("Demikian permohonan ini saya ajukan, atas perhatiannya diucapkan terima kasih.")
doc.add_paragraph("")

# TTD
ttd_table = doc.add_table(rows=4, cols=2)
ttd_table.rows[0].cells[0].text = "Mengetahui,"
ttd_table.rows[0].cells[1].text = "Pemohon,"
ttd_table.rows[1].cells[0].text = "Atasan Langsung"
ttd_table.rows[1].cells[1].text = ""
ttd_table.rows[2].cells[0].text = ""
ttd_table.rows[2].cells[1].text = ""
ttd_table.rows[3].cells[0].text = "{{KABID_KASI}}"
ttd_table.rows[3].cells[1].text = "{{NAMA}}"

doc.add_paragraph("")

# ===== PEMISAH =====
doc.add_paragraph("─" * 60)
doc.add_paragraph("")

# ===== BAGIAN 2: SURAT CUTI RESMI =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SURAT CUTI")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Nomor: {{NO_SURAT}}")
run.bold = True

doc.add_paragraph("")
doc.add_paragraph("Kepala Sub. Bag Umum dan Kepegawaian")
doc.add_paragraph("Dinas Perhubungan Kota Bogor")
doc.add_paragraph("")

table2 = doc.add_table(rows=6, cols=2)
table2.style = "Table Grid"
data2 = [
    ("Nama", "{{NAMA}}"),
    ("NIP", "{{NIP}}"),
    ("Jabatan", "{{JABATAN}}"),
    ("Bidang/Seksi", "{{SEKSI}}"),
    ("Tanggal Cuti", "{{TANGGAL_CUTI}}"),
    ("Keperluan", "{{KEPERLUAN}}"),
]
for i, (label, placeholder) in enumerate(data2):
    table2.rows[i].cells[0].text = label
    table2.rows[i].cells[1].text = placeholder

doc.add_paragraph("")
doc.add_paragraph("Menyetujui pemberian cuti tersebut di atas.")
doc.add_paragraph("")

# TTD Kasubag
ttd2 = doc.add_table(rows=4, cols=2)
ttd2.rows[0].cells[0].text = ""
ttd2.rows[0].cells[1].text = "Ka. Sub. Bag Umum dan Kepegawaian,"
ttd2.rows[1].cells[0].text = ""
ttd2.rows[1].cells[1].text = ""
ttd2.rows[2].cells[0].text = ""
ttd2.rows[2].cells[1].text = ""
ttd2.rows[3].cells[0].text = ""
ttd2.rows[3].cells[1].text = "ETIN SUHARTINI, S.E"

p_nip = doc.add_paragraph()
p_nip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_nip.add_run("NIP: 19770726 201101 2 002")

doc.save("template_surat/template_cuti_pkwt.docx")
print("Template created: template_surat/template_cuti_pkwt.docx")
